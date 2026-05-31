# services/file_creator_service.py
"""
Create structured files: .docx, .xlsx, .csv, .pdf
User can provide data directly via chat.
Save location is COMPULSORY — always asked before creating.
"""

import os
import re
import json
from pathlib import Path
# from services.llm_service import _call_ollama
# from utils.config import OLLAMA_FAST_MODEL


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _get_desktop_path() -> str:
    user_profile = os.environ.get("USERPROFILE", "")
    for path in [
        os.path.join(user_profile, "Desktop"),
        os.path.join(user_profile, "OneDrive", "Desktop"),
    ]:
        if os.path.exists(path):
            return path
    return os.getcwd()


def _resolve_location(location: str) -> str | None:
    """Resolve user-provided location string to a valid folder path."""
    if not location:
        return None
    loc = location.strip().strip('"').strip("'")

    # Drive shorthand
    if len(loc) == 1 and loc.isalpha():
        return f"{loc.upper()}:\\"
    if len(loc) == 2 and loc[1] == ":":
        return f"{loc.upper()}\\"

    # Named folders
    user_profile = os.environ.get("USERPROFILE", "")
    folder_map = {
        "desktop":   os.path.join(user_profile, "Desktop"),
        "documents": os.path.join(user_profile, "Documents"),
        "downloads": os.path.join(user_profile, "Downloads"),
        "pictures":  os.path.join(user_profile, "Pictures"),
        "videos":    os.path.join(user_profile, "Videos"),
        "music":     os.path.join(user_profile, "Music"),
    }
    if loc.lower() in folder_map:
        return folder_map[loc.lower()]

    return loc


def _resolve_path(filename: str, save_path: str) -> Path:
    return Path(save_path) / filename


# ─────────────────────────────────────────────────────────────
# LLM DATA EXTRACTOR
# ─────────────────────────────────────────────────────────────

_DATA_EXTRACT_PROMPT = """Extract structured data from this user message for file creation.

Return ONLY JSON (no markdown, no explanation):
{{
  "filename": "filename with extension or null",
  "file_type": "docx" | "xlsx" | "csv" | "pdf" | null,
  "title": "document title or null",
  "headers": ["col1", "col2"],
  "rows": [["val1", "val2"]],
  "content": "plain text content for docx/pdf or null",
  "location": "save path or null"
}}

Rules:
- For docx/pdf: put text in "content", leave headers/rows empty
- For xlsx/csv: put tabular data in headers + rows
- Extract filename if mentioned (e.g., "report.xlsx")
- Parse rows from patterns like "John 20 A, Mary 22 B" or "row1: a b c | row2: d e f"

Examples:
"create students.csv columns Name Age Grade rows John 20 A, Mary 22 B" ->
{{"filename":"students.csv","file_type":"csv","title":null,"headers":["Name","Age","Grade"],"rows":[["John","20","A"],["Mary","22","B"]],"content":null,"location":null}}

"make notes.docx with content Hello World" ->
{{"filename":"notes.docx","file_type":"docx","title":"notes","headers":[],"rows":[],"content":"Hello World","location":null}}

"create budget.xlsx columns Month Income Expense rows Jan 5000 3000, Feb 6000 4000" ->
{{"filename":"budget.xlsx","file_type":"xlsx","title":"Budget","headers":["Month","Income","Expense"],"rows":[["Jan","5000","3000"],["Feb","6000","4000"]],"content":null,"location":null}}

User message: "{message}"
JSON:"""


def parse_file_creation_intent(text: str) -> dict:
    """Pure regex-based parser (NO LLM)."""

    t = text.lower()

    result = {
        "filename": None,
        "file_type": None,
        "title": None,
        "headers": [],
        "rows": [],
        "content": None,
        "location": None
    }

    # Detect file type
    type_map = {
        "txt": ["txt", "notepad"],
        "docx": ["docx", "word", "word document"],
        "xlsx": ["xlsx", "excel", "spreadsheet"],
        "csv": ["csv"],
        "pdf": ["pdf"]
    }

    for ftype, keywords in type_map.items():
        if any(k in t for k in keywords):
            result["file_type"] = ftype
            break

    # Extract filename
    m = re.search(r"([\w\-]+\.(docx|xlsx|csv|pdf|txt))", text, re.I)
    if m:
        result["filename"] = m.group(1)

    # Extract headers
    col_match = re.search(r"columns?\s+(.+?)(?:rows?|$)", t)
    if col_match:
        result["headers"] = [
            x.strip().capitalize()
            for x in re.split(r"[,\s]+", col_match.group(1))
            if x.strip()
        ]

    # Extract rows
    row_match = re.search(r"rows?\s+(.+)", t)
    if row_match:
        raw_rows = row_match.group(1).split(",")
        for r in raw_rows:
            row = [x.strip() for x in r.split() if x.strip()]
            if row:
                result["rows"].append(row)

    # Extract content (for docx/pdf)
    content_match = re.search(r"(?:content|text|with)\s+(.+)", text, re.I)
    if content_match:
        result["content"] = content_match.group(1).strip()

    return result


def _regex_parse_creation(text: str) -> dict:
    t = text.lower()
    default = {
        "filename": None, "file_type": None, "title": None,
        "headers": [], "rows": [], "content": None, "location": None
    }
    for ext in ("docx", "xlsx", "csv", "pdf", "txt"):
        if ext in t or {"docx": "word", "xlsx": "excel", "csv": "csv", "pdf": "pdf", "txt": "txt"}[ext] in t:
            default["file_type"] = ext
            break
    m = re.search(r"[\w\-]+\.(docx|xlsx|csv|pdf|txt)", t)
    if m:
        default["filename"] = m.group(0)
    col_m = re.search(r"columns?\s+(.+?)(?:\s+(?:and\s+)?rows?|\s*$)", t)
    if col_m:
        default["headers"] = [h.strip() for h in re.split(r"[,\s]+", col_m.group(1)) if h.strip()]
    content_m = re.search(r"(?:content|text|with)\s+(.+?)(?:\s+save|\s+in|$)", t)
    if content_m:
        default["content"] = content_m.group(1).strip()
    return default


def is_file_creation_request(text: str) -> bool:
    t = text.lower()
    has_create = bool(re.search(r"\b(create|make|generate|new)\b", t))
    has_type = bool(re.search(r"\b(txt|docx|xlsx|csv|pdf|notepad|word doc|word document|excel|spreadsheet)\b", t))

    # If the user explicitly names multiple files, treat this as a general file operation
    # request rather than the single-file structured creator.
    raw_filenames = re.findall(r"\b[\w\-. ]+?\.(?:docx|xlsx|csv|pdf|txt)\b", text, re.I)
    filenames = []
    for raw in raw_filenames:
        raw = raw.strip()
        clean = re.sub(r"^(?:create|make|generate|new|and|or|file|document|the|a)\s+", "", raw, flags=re.I).strip()
        if clean:
            filenames.append(clean)
    if len(filenames) > 1:
        return False

    return has_create and has_type


# ─────────────────────────────────────────────────────────────
# CSV CREATOR
# ─────────────────────────────────────────────────────────────

def create_csv(filename: str, headers: list, rows: list, save_path: str) -> dict:
    try:
        import csv
        path = _resolve_path(filename, save_path)
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if headers:
                writer.writerow(headers)
            for row in rows:
                writer.writerow(row)
        return {
            "status": "success",
            "message": (
                f"✅ CSV file created!\n\n"
                f"📄 File: {filename}\n"
                f"📂 Location: {path.parent}\n"
                f"📊 Rows: {len(rows)}\n"
                f"📋 Columns: {', '.join(headers) if headers else 'none'}"
            ),
            "path": str(path)
        }
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to create CSV: {str(e)}"}


# ─────────────────────────────────────────────────────────────
# XLSX CREATOR
# ─────────────────────────────────────────────────────────────

def create_xlsx(filename: str, headers: list, rows: list, title: str = None, save_path: str = None) -> dict:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return {"status": "error", "message": "❌ openpyxl not installed.\nRun: pip install openpyxl"}

    try:
        path = _resolve_path(filename, save_path)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title or "Sheet1"

        header_fill = PatternFill(start_color="6366F1", end_color="6366F1", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)

        if headers:
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                ws.column_dimensions[cell.column_letter].width = max(15, len(str(header)) + 4)

        for row_idx, row in enumerate(rows, 2 if headers else 1):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        wb.save(path)
        return {
            "status": "success",
            "message": (
                f"✅ Excel file created!\n\n"
                f"📄 File: {filename}\n"
                f"📂 Location: {path.parent}\n"
                f"📊 Rows: {len(rows)}\n"
                f"📋 Columns: {', '.join(headers) if headers else 'none'}"
            ),
            "path": str(path)
        }
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to create Excel: {str(e)}"}


# ─────────────────────────────────────────────────────────────
# DOCX CREATOR
# ─────────────────────────────────────────────────────────────

def create_docx(filename: str, content: str = None, title: str = None,
                headers: list = None, rows: list = None, save_path: str = None) -> dict:
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"status": "error", "message": "❌ python-docx not installed.\nRun: pip install python-docx"}

    try:
        path = _resolve_path(filename, save_path)
        doc = Document()

        if title:
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content:
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        if headers and rows:
            doc.add_heading("Data", level=2)
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            hrow = table.rows[0]
            for i, header in enumerate(headers):
                cell = hrow.cells[i]
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
            for ri, row in enumerate(rows, 1):
                trow = table.rows[ri]
                for ci, val in enumerate(row):
                    if ci < len(headers):
                        trow.cells[ci].text = str(val)

        doc.save(path)
        return {
            "status": "success",
            "message": (
                f"✅ Word document created!\n\n"
                f"📄 File: {filename}\n"
                f"📂 Location: {path.parent}"
            ),
            "path": str(path)
        }
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to create DOCX: {str(e)}"}


# ─────────────────────────────────────────────────────────────
# PDF CREATOR
# ─────────────────────────────────────────────────────────────

def create_pdf(filename: str, content: str = None, title: str = None,
               headers: list = None, rows: list = None, save_path: str = None) -> dict:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError:
        return {"status": "error", "message": "❌ reportlab not installed.\nRun: pip install reportlab"}

    try:
        path = _resolve_path(filename, save_path)
        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        if title:
            title_style = ParagraphStyle(
                'CT', parent=styles['Title'],
                textColor=colors.HexColor('#6366F1'), fontSize=20, spaceAfter=20
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.5*cm))

        if content:
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 0.3*cm))

        if headers and rows:
            story.append(Spacer(1, 0.5*cm))
            table_data = [headers] + rows
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366F1')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(t)

        if not story:
            story.append(Paragraph("Empty document", styles['Normal']))

        doc.build(story)
        return {
            "status": "success",
            "message": (
                f"✅ PDF created!\n\n"
                f"📄 File: {filename}\n"
                f"📂 Location: {path.parent}"
            ),
            "path": str(path)
        }
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to create PDF: {str(e)}"}

def create_txt(filename: str, content: str = "", save_path: str = None) -> dict:
    try:
        path = _resolve_path(filename, save_path)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content or "")
        return {
            "status": "success",
            "message": (
                f"✅ Text file created!\n\n"
                f"📄 File: {filename}\n"
                f"📂 Location: {path.parent}"
            ),
            "path": str(path)
        }
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to create TXT: {str(e)}"}
    
# ─────────────────────────────────────────────────────────────
# MAIN HANDLER — ALWAYS ASKS LOCATION FIRST
# ─────────────────────────────────────────────────────────────

def handle_file_creation(user_prompt: str) -> dict:
    """
    Entry point for file creation.
    ALWAYS asks for save location before creating.
    """
    intent = parse_file_creation_intent(user_prompt)

    file_type = intent.get("file_type")
    filename = intent.get("filename")
    headers = intent.get("headers", [])
    rows = intent.get("rows", [])
    content = intent.get("content")
    title = intent.get("title")

    # Ask file type if missing
    if not file_type:
        return {
            "status": "need_type",
            "message": (
                "📄 What type of file do you want to create?\n\n"
                "  1. Word Document (.docx)\n"
                "  2. Excel Spreadsheet (.xlsx)\n"
                "  3. CSV File (.csv)\n"
                "  4. PDF Document (.pdf)\n"
                "  5. Text Document (.txt)\n\n"
                "Type the number or extension (e.g. xlsx):"
            ),
            "pending": {
                "action": "create_file",
                "filename": filename,
                "headers": headers,
                "rows": rows,
                "content": content,
                "title": title
            }
        }

    # Build default filename if not given
    if not filename:
        ext_map = {"docx": "document.docx", "xlsx": "spreadsheet.xlsx",
                   "csv": "data.csv", "pdf": "document.pdf", "txt": "document.txt"}
        filename = ext_map.get(file_type, f"file.{file_type}")

    # Ensure correct extension
    if not filename.lower().endswith(f".{file_type}"):
        filename = f"{Path(filename).stem}.{file_type}"

    # ✅ COMPULSORY: Always ask save location
    return {
        "status": "need_save_location",
        "message": (
            f"📂 Where should I save '{filename}'?\n\n"
            "Options:\n"
            "  1. Desktop (default)\n"
            "  2. Documents\n"
            "  3. Downloads\n"
            "  4. Custom path (e.g., D:\\\\Projects)\n\n"
            "Type 1, 2, 3, or a full path:"
        ),
        "pending": {
            "action": "create_file",
            "file_type": file_type,
            "filename": filename,
            "headers": headers,
            "rows": rows,
            "content": content,
            "title": title
        }
    }


def create_file_at_location(pending: dict, location_input: str) -> dict:
    """
    Called after user provides save location.
    Resolves the path and creates the file.
    """
    from services.file_advanced_service import _normalise_location

    file_type = pending.get("file_type")
    filename = pending.get("filename")
    headers = pending.get("headers", [])
    rows = pending.get("rows", [])
    content = pending.get("content")
    title = pending.get("title")

    # Resolve location
    loc_input = location_input.strip()

    # Numeric shortcuts
    user_profile = os.environ.get("USERPROFILE", "")
    shortcut_map = {
        "1": os.path.join(user_profile, "Desktop"),
        "desktop": os.path.join(user_profile, "Desktop"),
        "2": os.path.join(user_profile, "Documents"),
        "documents": os.path.join(user_profile, "Documents"),
        "3": os.path.join(user_profile, "Downloads"),
        "downloads": os.path.join(user_profile, "Downloads"),
    }

    if loc_input.lower() in shortcut_map:
        save_path = shortcut_map[loc_input.lower()]
    else:
        save_path = _normalise_location(loc_input) or loc_input

    if not os.path.exists(save_path):
        return {
            "status": "error",
            "message": (
                f"❌ Location not found: {save_path}\n\n"
                "Please enter a valid path, or type:\n"
                "  1 for Desktop  |  2 for Documents  |  3 for Downloads"
            )
        }

    if not os.path.isdir(save_path):
        return {"status": "error", "message": f"❌ '{save_path}' is not a folder."}

    path = _resolve_path(filename, save_path)
    if path.exists():
        return {
            "status": "exists",
            "message": (
                f"⚠️ File already exists: {filename}\n"
                f"📂 Location: {path.parent}\n\n"
                "Type 'yes' to overwrite or 'no' to cancel."
            ),
            "filepath": str(path),
            "filename": filename,
            "save_path": save_path,
            "file_type": file_type,
            "headers": headers,
            "rows": rows,
            "content": content,
            "title": title,
        }

    # Create the file
    if file_type == "csv":
        return create_csv(filename, headers, rows, save_path)
    elif file_type == "xlsx":
        return create_xlsx(filename, headers, rows, title, save_path)
    elif file_type == "docx":
        return create_docx(filename, content, title,
                           headers or None, rows or None, save_path)
    elif file_type == "pdf":
        return create_pdf(filename, content, title,
                          headers or None, rows or None, save_path)
    elif file_type == "txt":
        return create_txt(filename, content, save_path)
    else:
        return {"status": "error", "message": f"❌ Unsupported file type: {file_type}"}


def add_data_to_file(filename: str, rows: list, file_path: str = None) -> dict:
    """Append data rows to existing xlsx or csv."""
    if not file_path:
        desktop = _get_desktop_path()
        file_path = os.path.join(desktop, filename)

    path = Path(file_path)
    if not path.exists():
        return {"status": "error", "message": f"❌ File not found: {filename}"}

    ext = path.suffix.lower()
    try:
        if ext == ".csv":
            import csv
            with open(path, 'a', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for row in rows:
                    writer.writerow(row)
            return {"status": "success", "message": f"✅ Added {len(rows)} row(s) to {filename}"}
        elif ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            for row in rows:
                ws.append(row)
            wb.save(path)
            return {"status": "success", "message": f"✅ Added {len(rows)} row(s) to {filename}"}
        else:
            return {"status": "error", "message": f"❌ Cannot add rows to {ext} files."}
    except Exception as e:
        return {"status": "error", "message": f"❌ Failed to add data: {str(e)}"}